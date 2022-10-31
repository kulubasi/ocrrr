from django.shortcuts import render, redirect
from django.http import HttpResponse,JsonResponse
from django.views.generic import FormView
from .forms import *
from django.views.decorators.csrf import csrf_exempt
from django.core.files.storage import FileSystemStorage
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.db import IntegrityError

import pytesseract
from .forms import uploadform
import os, sys
from PIL import Image
from django.conf import settings
#from .models import Book

import docx
import PyPDF2
#from xpdf_python import to_text
import pdfplumber
from pdf2image import convert_from_path
import re
import traceback




# Create your views here.



def home(request):


    return render(request,"home.html")

def login_user(request):
    if request.method=="POST":
        username=request.POST.get('username')
        password=request.POST.get('password1')
        psswd2=request.POST.get('password2')
        user=authenticate(request,username=username,password=password)
        
        if password != psswd2:
            return render(request,"login.html",{'form':UserCreationForm,'Error':'The Passwords are not matching!'})
        if user is not None:
            login(request,user)
            if request.user.is_authenticated:
                Username = {request.user.username}
                return redirect('home')
        else:
            messages.success(request,("There was an error Logging in Please Try Aagain"))
            return redirect('login')


    else:
        print("User is not logged in :(")

    return render(request,"login.html",{'form':UserCreationForm})

def logout_user(request):
    logout(request)
    return redirect('home')

def profile(request):
    return render(request,"profile.html")



def signup(request):
    if request.method=="POST":
        if request.POST.get('password1')==request.POST.get('password2'):# password1 and password2 are from the user cretion form after inspecting source
            try:
                saveuser=User.objects.create_user(request.POST.get('username'), password=request.POST.get('password1'))
                saveuser.save()
                return render(request,"signup.html",{'form':UserCreationForm,'Info':'The user '+request.POST.get('username')+' is saved successfully...!'})
            

            except IntegrityError:#this is for a user who already exixts
                return render(request,"signup.html",{'form':UserCreationForm,'Error':'The user '+request.POST.get('username')+' already exists...!'})

        else:
            return render(request,"signup.html",{'form':UserCreationForm,'Error':'The Passwords are not matching!'})   

    else:
        return render(request,"signup.html",{'form':UserCreationForm})

def upload(request):
    global text
    text=""
    message=""
    msg=""
    form=uploadform()
    # form=uploadform(request.POST,request.FILES)
    if request.method=='POST':
        # stimg=request.FILES['img']
        # fs=FileSystemStorage()
        # fs.save(stimg.name,stimg)
        try:
            form=uploadform(request.POST,request.FILES)
            if form.is_valid():
                form.save()
                stimg=request.FILES['img']
                stimg=stimg.name
                print(stimg)
                path=settings.MEDIA_ROOT
                pathz=path+"/imagegallery/"+stimg
                pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
                text=pytesseract.image_to_string(Image.open(pathz), lang='eng')
                os.remove(pathz)
                return redirect("ready")
        except Exception:
            message=traceback.print_exc()
    else:
        messager="Ebintu eeh"
    context={'form':form,'text':text,'messg':message,'messgr':msg}
    return render(request,'new.html',context)


def uploadpdf(request):
    global text
    text=""
    message=""
    msg=""
    form=uploadpdfform()
    if request.method=='POST':
        form=uploadpdfform(request.POST,request.FILES)
        if form.is_valid():
            form.save()
            stimgi=request.FILES['pdf']
            stimg=stimgi.name
            
            remove = " "
            for char in remove:
                stimg = stimg.replace(char,"_")
            removed = " !()[]"
            for char in removed:
                stimg = stimg.replace(char,"")
            print(stimg)
            print("yes/n yes /n yesss")
            path=settings.MEDIA_ROOT
            pathy=os.path.join(path,'pdfs')
            pathz=os.path.join(pathy,stimg)
            print(pathz)
            # pathz=path.append("/pdfs"+stimg)
            # pdfreader=PyPDF2.PdfFileReader(pathz)
            # pageobj=pdfreader.getPage(0)
            # text=pageobj.extractText()

            pdf=pdfplumber.open(pathz)
            for page in pdf.pages:
                # text.append(page.extract_text())
                text=text+page.extract_text()+'~~~pgend~~~'

            return redirect("pdfready")


        # try:
        #     form=uploadpdfform(request.POST,request.FILES)
        #     if form.is_valid():
        #         form.save()
        #         stimg=request.FILES['pdf']
        #         stimg=stimg.name
        #         print(stimg)
        #         print("yes/n yes /n yesss")
        #         path=settings.MEDIA_ROOT
        #         pathz=path+"/pdfs/"+stimg
        #         # pdfreader=PyPDF2.PdfFileReader(pathz)
        #         # pageobj=pdfreader.getPage(0)
        #         # text=pageobj.extractText()

        #         pdf=pdfplumber.open(pathz)
        #         for page in pdf.pages:
        #             text.append(page.extract_text())
        #         os.remove(pathz)
        #         return redirect("ready")
        # except:
        #     message="Things arent yet right"
    else:
        messager="Ebintu eeh"
    context={'form':form,'text':text,'messg':message,'messgr':msg}
    return render(request,'pdfupload.html',context)



def pdf2img2txt(request):
    global text
    text=""
    message=""
    msg=""
    form=uploadpdfform()
    if request.method=='POST':
        form=uploadpdfform(request.POST,request.FILES)
        if form.is_valid():
            form.save()
            stimg=request.FILES['pdf']
            stimg=stimg.name
            remove = " "
            for char in remove:
                stimg = stimg.replace(char,"_")
            removed = " !()[]"
            for char in removed:
                stimg = stimg.replace(char,"")
            print(stimg)
            print("yes/n yes /n yesss")
            path=settings.MEDIA_ROOT
            pathy=os.path.join(path,'pdfs')
            pathz=os.path.join(pathy,stimg)
            popplerpath=os.path.join(settings.BASE_DIR,'poppler-0.68.0','bin')

            
            pages = convert_from_path(pathz,500,poppler_path=popplerpath)
            i = 1
            print(popplerpath)
            createdimages=[]  #images directory stored here
            for page in pages:
                mine = "Page_" + str(i) + str(stimg)+".jpg"
                # page.save(r"D:\Project\global.jpg", "JPEG")
                page.save(os.path.join(pathy,mine),"JPEG")
                createdimages.append(mine)
                i = i+1
                print(mine)

            for imagepath in createdimages:
                newpath=os.path.join(pathy,imagepath)
                pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
                text+=pytesseract.image_to_string(Image.open(newpath), lang='eng')+'~~~pgend~~~'


            return redirect("pdfready")


    else:
        messager="Ebintu eeh"
    context={'form':form,'text':text,'messg':message,'messgr':msg}
    return render(request,'pdfupload.html',context)



def ready(request):
    #h=input("Enter your name")
    h=5
    #return HttpResponse("Things are gonna be okay after lunch {}".format(h))
    context={'text':text}
    return render(request, 'ready.html',context)

def pdfready(request):
    #h=input("Enter your name")
    h=5
    #return HttpResponse("Things are gonna be okay after lunch {}".format(h))
    context={'text':text}
    return render(request, 'pdfready.html',context)

def download(request):
    mydc=docx.Document()
    mydc.add_paragraph(text)
    response=HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition']='attachment; filename=download.docx'

    mydc.save(response)
    return response

    # context={'text':text}
    # return render(request,'download.html',context)

def pdfdownload(request):  #downloads word doc for pdfs
    mydc=docx.Document()
    # mydc.add_paragraph(text)

    # document = Document()
    #
    # document.add_heading('Document Title', 0)
    #
    # document.add_paragraph('Intense quote', style='IntenseQuote')
    #
    # document.add_paragraph(
    # 'first item in unordered list', style='ListBullet'
    # )
    # splitted_list=text.split('\n')
    for char in text.split('~~~pgend~~~'):
        mydc.add_paragraph(char)
        mydc.add_page_break()
    response=HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition']='attachment; filename=download.docx'

    mydc.save(response)
    return response

    # context={'text':text}
    # return render(request,'download.html',context)


def uploadi(request):
    # form= uploadform({})
    #form= uploadform(request.POST,request.FILES)
    form= uploadform(request.POST or request.FILES)
    #
    #global myimg
    # myimg=''
    # context={'form':form,'k':myimg}
    context={'form':form}
    if form.is_valid():
        form.save()
        # myimg= form.cleaned_data.get("first_name")
        #myim= form.cleaned_data.get("imge")
        myim=request.FILES['imge']
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        #print(form.is_bound)
        myimg = pytesseract.image_to_string(myim, lang='eng')
        context={'k':myimg}
    else:
        myimg="upload request failed"

    # if request.method == 'POST'and form.is_valid=="True":
    #     submitbutton= request.POST.get("submit")
    #     t="am request Post is passed"
    #     #if form.is_valid():
    #     p="form validation is passed"
    #     form= uploadform(request.POST or request.FILES)
    #     first_name= form.cleaned_data.get("first_name")
    #     myimg= form.cleaned_data.get("image")
    #     print(first_name)
    #     pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    #     #mym=Book._meta.get_field('cover')
    #     text = pytesseract.image_to_string(myimg, lang='eng')
    #     #redirect('scan')
    #     context={'form': form, 'firstname': first_name,'imge':myimg, 'k':t,'h':text}

    return render(request,'new.html',context)

# def scan(request):
#     if request.method == 'POST'and form.is_valid=="True":
#         submitbutton= request.POST.get("submit")
#         t="am request Post is passed"
#         #if form.is_valid():
#         p="form validation is passed"
#         form= uploadform(request.POST or request.FILES)
#         first_name= form.cleaned_data.get("first_name")
#         myimg= form.cleaned_data.get("image")
#         print(first_name)
#         pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
#         #mym=Book._meta.get_field('cover')
#         text = pytesseract.image_to_string(myimg, lang='eng')
#         #redirect('scan')
#         context={'form': form, 'firstname': first_name,'imge':myimg, 'k':t,'h':text}
#     return render(request,'scan.html',context)

def book_list(request):
    books=Book.objects.all()
    return render(request,'book_list.html', {'books':books})

# def upload_book(request):
#     if request.method == 'POST':
#         form=BookForm(request.POST,request.FILES)
#         if form.is_valid():
#             form.save()
#             return redirect('book_list')
#     else:
#         form=BookForm()
#     return render(request,'upload_book.html',{'form':form})


# WORKING
# def upload(request):
#     text=""
#     message=""
#     msg=""
#     # form=uploadform(request.POST,request.FILES)
#     if request.method=='POST':
#         # stimg=request.FILES['img']
#         # fs=FileSystemStorage()
#         # fs.save(stimg.name,stimg)
#         try:
#             form=uploadform(request.POST,request.FILES)
#             if form.is_valid():
#                 form.save()
#                 stimg=request.FILES['img']
#                 stimg=stimg.name
#                 print(stimg)
#                 path=settings.MEDIA_ROOT
#                 pathz=path+"/imagegallery/"+stimg
#                 pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
#                 text=pytesseract.image_to_string(Image.open(pathz), lang='eng')
#                 os.remove(pathz)
#         except:
#             message="Things arent yet right"
#     else:
#         messager="Ebintu eeh"
#     context={'form':form,'text':text,'messg':message,'messgr':msg}
#     return render(request,'new.html',context)
